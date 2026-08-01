"""文档解析器（Phase 7.2 需求一/十一）。

原生零依赖支持：txt / md / json / csv / tsv / html
可选依赖（缺失自动降级并给出可读提示，绝不抛异常）：
  - pdf   → pypdf / PyPDF2
  - xlsx  → openpyxl
  - docx  → python-docx
"""
from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field

from backend.multimodal.text.extractor import Entity, detect_subtype, extract_entities, summarize

_TAG_RE = re.compile(r"<[^>]+>")


@dataclass
class DocumentResult:
    ok: bool = False
    kind: str = ""          # txt/csv/json/pdf/xlsx/docx/html
    text: str = ""
    rows: int = 0
    subtype: str = ""
    entities: list[Entity] = field(default_factory=list)
    summary: str = ""
    message: str = ""

    def to_dict(self) -> dict:
        return {
            "ok": self.ok,
            "kind": self.kind,
            "rows": self.rows,
            "subtype": self.subtype,
            "entities": [e.to_dict() for e in self.entities],
            "summary": self.summary,
            "message": self.message,
        }


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _parse_csv(data: bytes, delimiter: str = ",") -> tuple[str, int]:
    text = _decode(data)
    try:
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        lines = [" | ".join(str(c).strip() for c in row if str(c).strip()) for row in reader]
    except csv.Error:
        lines = text.split("\n")
    lines = [ln for ln in lines if ln]
    return "\n".join(lines[:2000]), len(lines)


def _parse_json(data: bytes) -> tuple[str, int]:
    text = _decode(data)
    try:
        obj = json.loads(text)
    except json.JSONDecodeError:
        return text[:20000], 0

    lines: list[str] = []

    def walk(node, prefix: str = "") -> None:
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, f"{prefix}{k}: " if not isinstance(v, (dict, list)) else prefix)
        elif isinstance(node, list):
            for item in node:
                walk(item, prefix)
        else:
            lines.append(f"{prefix}{node}")

    walk(obj)
    return "\n".join(lines[:2000]), len(lines)


def _parse_pdf(data: bytes) -> tuple[str, str]:
    try:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:  # noqa: BLE001
            from PyPDF2 import PdfReader  # type: ignore
    except Exception:  # noqa: BLE001
        return "", "服务器未安装 PDF 解析组件（pypdf），请改上传 CSV/Excel 或截图。"
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(p.extract_text() or "") for p in reader.pages[:50]]
        return "\n".join(pages), ""
    except Exception:  # noqa: BLE001
        return "", "PDF 解析失败，文件可能是扫描件或已加密，可改用截图上传。"


def _parse_xlsx(data: bytes) -> tuple[str, int, str]:
    try:
        import openpyxl  # type: ignore
    except Exception:  # noqa: BLE001
        return "", 0, "服务器未安装 Excel 解析组件（openpyxl），请另存为 CSV 后上传。"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# 工作表 {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
                if len(lines) > 2000:
                    break
        return "\n".join(lines), len(lines), ""
    except Exception:  # noqa: BLE001
        return "", 0, "Excel 解析失败，请检查文件是否损坏。"


def _parse_docx(data: bytes) -> tuple[str, str]:
    try:
        import docx  # type: ignore
    except Exception:  # noqa: BLE001
        return "", "服务器未安装 Word 解析组件（python-docx），请另存为 PDF 或 TXT。"
    try:
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts), ""
    except Exception:  # noqa: BLE001
        return "", "Word 解析失败，请检查文件是否损坏。"


def parse_document(filename: str, data: bytes) -> DocumentResult:
    """解析文件 → 文本 → 财富实体（绝不写库）。"""
    name = (filename or "").lower()
    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    text, rows, kind, message = "", 0, ext.lstrip(".") or "txt", ""

    if ext == ".csv":
        text, rows = _parse_csv(data)
    elif ext == ".tsv":
        text, rows = _parse_csv(data, delimiter="\t")
    elif ext == ".json":
        text, rows = _parse_json(data)
    elif ext == ".pdf":
        text, message = _parse_pdf(data)
    elif ext in (".xlsx", ".xls"):
        text, rows, message = _parse_xlsx(data)
    elif ext in (".docx", ".doc"):
        text, message = _parse_docx(data)
    elif ext in (".html", ".htm"):
        text = _TAG_RE.sub(" ", _decode(data))
    else:
        text = _decode(data)
        kind = "txt"

    text = (text or "").strip()
    if not text:
        return DocumentResult(
            ok=False, kind=kind,
            message=message or "未能从文件中读到文本内容，请确认文件格式。",
        )

    entities = extract_entities(text)
    return DocumentResult(
        ok=bool(entities),
        kind=kind,
        text=text[:20000],
        rows=rows,
        subtype=detect_subtype(text),
        entities=entities,
        summary=summarize(text, entities),
        message=message or ("" if entities else "文件已解析，但未识别到明确的金额数据。"),
    )
